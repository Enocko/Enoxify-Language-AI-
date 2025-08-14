import os
import tempfile
import time
from typing import Dict, Any, Optional
import PyPDF2
from docx import Document
import openai
from PIL import Image
import pytesseract

class DocumentProcessor:
    def __init__(self):
        self.openai_api_key = os.getenv("OPENAI_API_KEY")
        self.client = openai.OpenAI(api_key=self.openai_api_key) if self.openai_api_key else None
        
        # Supported file formats
        self.supported_formats = {
            "pdf": [".pdf"],
            "word": [".docx", ".doc"],
            "text": [".txt", ".md"],
            "image": [".png", ".jpg", ".jpeg", ".gif", ".bmp"],
            "audio": [".mp3", ".wav", ".m4a", ".flac"],
            "video": [".mp4", ".avi", ".mov", ".mkv"]
        }
    
    async def process(self, file_path: str, output_format: str = "both", 
                     include_audio: bool = True, include_simplified_text: bool = True) -> Dict[str, Any]:
        """
        Process various document formats and convert to accessible formats
        
        Args:
            file_path: Path to the document file
            output_format: Desired output format (audio, simplified_text, both)
            include_audio: Whether to include audio output
            include_simplified_text: Whether to include simplified text output
            
        Returns:
            Dictionary with processing results
        """
        start_time = time.time()
        
        try:
            # Determine file type
            file_extension = os.path.splitext(file_path)[1].lower()
            file_type = self._get_file_type(file_extension)
            
            if not file_type:
                raise Exception(f"Unsupported file format: {file_extension}")
            
            # Extract text content based on file type
            extracted_text = await self._extract_text(file_path, file_type)
            
            if not extracted_text:
                raise Exception("Could not extract text from document")
            
            # Process the extracted text
            result = {
                "original_format": file_type,
                "output_formats": [],
                "processing_time": time.time() - start_time
            }
            
            # Generate simplified text if requested
            if include_simplified_text:
                from .text_simplifier import TextSimplifier
                simplifier = TextSimplifier()
                simplified_result = await simplifier.simplify(extracted_text, "middle_school", True)
                
                # Handle new simplifier return format
                if isinstance(simplified_result, dict):
                    result["simplified_text"] = simplified_result["simplified_text"]
                else:
                    result["simplified_text"] = simplified_result
                    
                result["output_formats"].append("simplified_text")
            
            # Generate audio if requested
            if include_audio:
                from .text_to_speech import TextToSpeech
                tts = TextToSpeech()
                tts_result = await tts.convert(extracted_text, "neutral", 1.0, "en-US")
                
                # Handle new TTS return format - now returns just filename
                if isinstance(tts_result, dict):
                    # TTS service now saves directly to temp directory and returns filename
                    result["audio_file_path"] = tts_result["audio_file_path"]
                else:
                    # Fallback for old format
                    result["audio_file_path"] = tts_result
                    
                result["output_formats"].append("audio")
            
            # Add file size information
            result["file_size"] = os.path.getsize(file_path)
            
            return result
            
        except Exception as e:
            raise Exception(f"Error processing document: {str(e)}")
    
    def _get_file_type(self, file_extension: str) -> Optional[str]:
        """Determine file type based on extension"""
        for file_type, extensions in self.supported_formats.items():
            if file_extension in extensions:
                return file_type
        return None
    
    async def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text content from different file types"""
        try:
            if file_type == "pdf":
                return await self._extract_from_pdf(file_path)
            elif file_type == "word":
                return await self._extract_from_word(file_path)
            elif file_type == "text":
                return await self._extract_from_text(file_path)
            elif file_type == "image":
                return await self._extract_from_image(file_path)
            elif file_type == "audio":
                return await self._extract_from_audio(file_path)
            elif file_type == "video":
                return await self._extract_from_video(file_path)
            else:
                raise Exception(f"Unsupported file type: {file_type}")
        except Exception as e:
            raise Exception(f"Error extracting text from {file_type}: {str(e)}")
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            text = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"PDF extraction error: {str(e)}")
    
    async def _extract_from_word(self, file_path: str) -> str:
        """Extract text from Word documents"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Word document extraction error: {str(e)}")
    
    async def _extract_from_text(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read().strip()
                except UnicodeDecodeError:
                    continue
            raise Exception("Could not decode text file with any supported encoding")
        except Exception as e:
            raise Exception(f"Text file extraction error: {str(e)}")
    
    async def _extract_from_image(self, file_path: str) -> str:
        """Extract text from images using OCR"""
        try:
            # Use pytesseract for OCR
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image)
            
            if not text.strip():
                # If OCR fails, try using OpenAI's Vision API
                if self.client:
                    text = await self._extract_with_openai_vision(file_path)
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Image OCR error: {str(e)}")
    
    async def _extract_with_openai_vision(self, file_path: str) -> str:
        """Extract text from image using OpenAI Vision API"""
        try:
            with open(file_path, "rb") as image_file:
                response = self.client.chat.completions.create(
                    model="gpt-4-vision-preview",
                    messages=[
                        {
                            "role": "user",
                            "content": [
                                {"type": "text", "text": "Extract all the text from this image. Return only the text content without any additional formatting or explanations."},
                                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_file.read()}"}}
                            ]
                        }
                    ],
                    max_tokens=1000
                )
            
            return response.choices[0].message.content.strip()
        except Exception as e:
            raise Exception(f"OpenAI Vision API error: {str(e)}")
    
    async def _extract_from_audio(self, file_path: str) -> str:
        """Extract text from audio files"""
        try:
            from .speech_to_text import SpeechToText
            stt = SpeechToText()
            
            # Create a mock file object for the speech-to-text service
            class MockFile:
                def __init__(self, path):
                    self.file = open(path, 'rb')
                    self.filename = path
                
                def read(self):
                    return self.file.read()
            
            mock_file = MockFile(file_path)
            result = await stt.transcribe(mock_file, "en-US", True)
            
            return result["text"]
        except Exception as e:
            raise Exception(f"Audio extraction error: {str(e)}")
    
    async def _extract_from_video(self, file_path: str) -> str:
        """Extract text from video files (audio + any text overlays)"""
        try:
            # For now, we'll extract audio and then transcribe it
            # In a production system, you might also want to extract text overlays
            from .speech_to_text import SpeechToText
            stt = SpeechToText()
            
            # Create a mock file object
            class MockFile:
                def __init__(self, path):
                    self.file = open(path, 'rb')
                    self.filename = path
                
                def read(self):
                    return self.file.read()
            
            mock_file = MockFile(file_path)
            result = await stt.transcribe(mock_file, "en-US", True)
            
            return result["text"]
        except Exception as e:
            raise Exception(f"Video extraction error: {str(e)}")
    
    def get_supported_formats(self) -> Dict[str, list]:
        """Get list of supported file formats"""
        return self.supported_formats
    
    def validate_file(self, file_path: str) -> Dict[str, Any]:
        """Validate if a file can be processed"""
        try:
            if not os.path.exists(file_path):
                return {"valid": False, "error": "File does not exist"}
            
            file_size = os.path.getsize(file_path)
            max_size = 100 * 1024 * 1024  # 100MB limit
            
            if file_size > max_size:
                return {"valid": False, "error": f"File too large ({file_size / 1024 / 1024:.1f}MB). Maximum size is 100MB."}
            
            file_extension = os.path.splitext(file_path)[1].lower()
            file_type = self._get_file_type(file_extension)
            
            if not file_type:
                return {"valid": False, "error": f"Unsupported file format: {file_extension}"}
            
            return {
                "valid": True,
                "file_type": file_type,
                "file_size": file_size,
                "supported": True
            }
            
        except Exception as e:
            return {"valid": False, "error": str(e)}
    
    def get_processing_options(self, file_type: str) -> Dict[str, Any]:
        """Get available processing options for a file type"""
        options = {
            "pdf": {
                "text_extraction": True,
                "simplification": True,
                "audio_conversion": True,
                "ocr_enhancement": False
            },
            "word": {
                "text_extraction": True,
                "simplification": True,
                "audio_conversion": True,
                "formatting_preservation": True
            },
            "text": {
                "text_extraction": True,
                "simplification": True,
                "audio_conversion": True,
                "language_detection": True
            },
            "image": {
                "text_extraction": True,
                "simplification": True,
                "audio_conversion": True,
                "ocr_enhancement": True
            },
            "audio": {
                "text_extraction": True,
                "simplification": True,
                "audio_conversion": False,
                "transcription_enhancement": True
            },
            "video": {
                "text_extraction": True,
                "simplification": True,
                "audio_conversion": False,
                "subtitle_generation": True
            }
        }
        
        return options.get(file_type, {})
